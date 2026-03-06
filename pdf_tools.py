import io
import os
import shutil
import subprocess
import tempfile
from typing import List, Optional, Tuple
from pypdf import PdfReader, PdfWriter
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from reportlab.lib.colors import Color, black
import fitz
from PIL import Image
import pytesseract
import config as app_config

def _bytesio(data: bytes) -> io.BytesIO:
    buf = io.BytesIO()
    buf.write(data)
    buf.seek(0)
    return buf

def merge_pdfs(buffers: List[io.BytesIO]) -> io.BytesIO:
    writer = PdfWriter()
    for b in buffers:
        reader = PdfReader(b)
        for p in reader.pages:
            writer.add_page(p)
    out = io.BytesIO()
    writer.write(out)
    out.seek(0)
    return out

def split_pdf_range(buffer: io.BytesIO, start: int, end: int) -> io.BytesIO:
    reader = PdfReader(buffer)
    writer = PdfWriter()
    n = len(reader.pages)
    s = max(1, start)
    e = min(n, end)
    for i in range(s - 1, e):
        writer.add_page(reader.pages[i])
    out = io.BytesIO()
    writer.write(out)
    out.seek(0)
    return out

def rotate_pdf(buffer: io.BytesIO, angle: int, pages: Optional[List[int]] = None) -> io.BytesIO:
    reader = PdfReader(buffer)
    writer = PdfWriter()
    indexes = None
    if pages:
        indexes = {p - 1 for p in pages if p > 0}
    for i, page in enumerate(reader.pages):
        if indexes is None or i in indexes:
            page.rotate(angle)
        writer.add_page(page)
    out = io.BytesIO()
    writer.write(out)
    out.seek(0)
    return out

def _make_text_watermark(text: str, opacity: float, font_size: int, width: float, height: float) -> io.BytesIO:
    pkt = io.BytesIO()
    c = canvas.Canvas(pkt, pagesize=(width, height))
    c.setFont("Helvetica", font_size)
    c.setFillColor(Color(0, 0, 0, alpha=opacity))
    c.drawCentredString(width / 2, height / 2, text)
    c.save()
    pkt.seek(0)
    return pkt

def watermark_pdf_text(buffer: io.BytesIO, text: str, opacity: float, font_size: int) -> io.BytesIO:
    reader = PdfReader(buffer)
    writer = PdfWriter()
    for page in reader.pages:
        media = page.mediabox
        w = float(media.width)
        h = float(media.height)
        wm_pdf = PdfReader(_make_text_watermark(text, opacity, font_size, w, h))
        wm_page = wm_pdf.pages[0]
        page.merge_page(wm_page)
        writer.add_page(page)
    out = io.BytesIO()
    writer.write(out)
    out.seek(0)
    return out

def add_page_numbers(buffer: io.BytesIO, position: str, font_size: int) -> io.BytesIO:
    reader = PdfReader(buffer)
    writer = PdfWriter()
    total = len(reader.pages)
    for idx, page in enumerate(reader.pages, start=1):
        media = page.mediabox
        w = float(media.width)
        h = float(media.height)
        pkt = io.BytesIO()
        c = canvas.Canvas(pkt, pagesize=(w, h))
        c.setFont("Helvetica", font_size)
        c.setFillColor(black)
        text = f"{idx}/{total}"
        if position == "top_left":
            c.drawString(20, h - 20, text)
        elif position == "top_right":
            c.drawRightString(w - 20, h - 20, text)
        elif position == "bottom_left":
            c.drawString(20, 20, text)
        elif position == "bottom_right":
            c.drawRightString(w - 20, 20, text)
        else:
            c.drawCentredString(w / 2, 20, text)
        c.save()
        pkt.seek(0)
        overlay = PdfReader(pkt).pages[0]
        page.merge_page(overlay)
        writer.add_page(page)
    out = io.BytesIO()
    writer.write(out)
    out.seek(0)
    return out

def protect_pdf(buffer: io.BytesIO, user_password: str, owner_password: str) -> io.BytesIO:
    reader = PdfReader(buffer)
    writer = PdfWriter()
    for p in reader.pages:
        writer.add_page(p)
    if not owner_password:
        owner_password = user_password
    writer.encrypt(user_password=user_password, owner_password=owner_password)
    out = io.BytesIO()
    writer.write(out)
    out.seek(0)
    return out

def unlock_pdf(buffer: io.BytesIO, password: str) -> io.BytesIO:
    reader = PdfReader(buffer, password=password)
    writer = PdfWriter()
    for p in reader.pages:
        writer.add_page(p)
    out = io.BytesIO()
    writer.write(out)
    out.seek(0)
    return out

def compress_pdf(buffer: io.BytesIO) -> io.BytesIO:
    try:
        return compress_pdf_ghostscript(buffer.getvalue(), "screen")
    except Exception:
        out = io.BytesIO()
        out.write(buffer.getvalue())
        out.seek(0)
        return out

def _which(names: List[str], config_key: Optional[str] = None) -> Optional[str]:
    if config_key:
        cfg = app_config.load()
        candidate = cfg.get(config_key)
        if candidate and os.path.isfile(candidate):
            return candidate
    
    for n in names:
        p = shutil.which(n)
        if p:
            return p
        
        # Check if the executable is directly in PATH (Windows issue)
        path_dirs = os.environ.get('PATH', '').split(os.pathsep)
        for path_dir in path_dirs:
            path_dir = path_dir.strip()
            if path_dir.endswith('.exe') and path_dir.lower().endswith(n.lower()):
                # PATH contains the full executable path
                if os.path.isfile(path_dir):
                    return path_dir
            else:
                # Normal directory check
                exe_path = os.path.join(path_dir, n)
                if os.path.isfile(exe_path):
                    return exe_path
                exe_path_ext = os.path.join(path_dir, n + '.exe')
                if os.path.isfile(exe_path_ext):
                    return exe_path_ext
    
    return None

def pdf_to_images(buffer: io.BytesIO, fmt: str = "png", dpi: int = 144) -> List[Tuple[str, bytes]]:
    doc = fitz.open(stream=buffer.getvalue(), filetype="pdf")
    out = []
    for i in range(len(doc)):
        page = doc.load_page(i)
        zoom = dpi / 72.0
        mat = fitz.Matrix(zoom, zoom)
        pix = page.get_pixmap(matrix=mat, alpha=False)
        data = None
        ext = fmt.lower()
        if ext == "png":
            data = pix.tobytes("png")
        elif ext in ("jpeg", "jpg"):
            data = pix.tobytes("jpg")
            ext = "jpg"
        else:
            data = pix.tobytes("png")
            ext = "png"
        out.append((f"page_{i+1}.{ext}", data))
    return out

def images_to_pdf(image_buffers: List[io.BytesIO]) -> io.BytesIO:
    pil_images = []
    for b in image_buffers:
        img = Image.open(b).convert("RGB")
        pil_images.append(img)
    first = pil_images[0]
    rest = pil_images[1:]
    out = io.BytesIO()
    first.save(out, format="PDF", save_all=True, append_images=rest)
    out.seek(0)
    return out

def _run_cmd(cmd: List[str], cwd: Optional[str] = None, timeout: int = 120) -> bytes:
    p = subprocess.run(cmd, cwd=cwd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=timeout)
    if p.returncode != 0:
        raise RuntimeError(p.stderr.decode("utf-8", "ignore"))
    return p.stdout

def docx_to_pdf(file_bytes: bytes) -> io.BytesIO:
    soffice = _which(["soffice", "soffice.exe"], "libreoffice_path")
    if not soffice:
        raise RuntimeError("LibreOffice not found. Please install LibreOffice from https://www.libreoffice.org/download/download-libreoffice/ and set the path in Settings, or ensure it's in your system PATH.")
    with tempfile.TemporaryDirectory() as td:
        in_path = os.path.join(td, "input.docx")
        out_dir = td
        with open(in_path, "wb") as f:
            f.write(file_bytes)
        _run_cmd([soffice, "--headless", "--convert-to", "pdf", "--outdir", out_dir, in_path])
        out_path = os.path.join(td, "input.pdf")
        with open(out_path, "rb") as f:
            data = f.read()
        return _bytesio(data)

def xlsx_to_pdf(file_bytes: bytes) -> io.BytesIO:
    soffice = _which(["soffice", "soffice.exe"], "libreoffice_path")
    if not soffice:
        raise RuntimeError("LibreOffice not found. Please install LibreOffice from https://www.libreoffice.org/download/download-libreoffice/ and set the path in Settings, or ensure it's in your system PATH.")
    with tempfile.TemporaryDirectory() as td:
        in_path = os.path.join(td, "input.xlsx")
        out_dir = td
        with open(in_path, "wb") as f:
            f.write(file_bytes)
        _run_cmd([soffice, "--headless", "--convert-to", "pdf", "--outdir", out_dir, in_path])
        out_path = os.path.join(td, "input.pdf")
        with open(out_path, "rb") as f:
            data = f.read()
        return _bytesio(data)

def pptx_to_pdf(file_bytes: bytes) -> io.BytesIO:
    soffice = _which(["soffice", "soffice.exe"], "libreoffice_path")
    if not soffice:
        raise RuntimeError("LibreOffice not found. Please install LibreOffice from https://www.libreoffice.org/download/download-libreoffice/ and set the path in Settings, or ensure it's in your system PATH.")
    with tempfile.TemporaryDirectory() as td:
        in_path = os.path.join(td, "input.pptx")
        out_dir = td
        with open(in_path, "wb") as f:
            f.write(file_bytes)
        _run_cmd([soffice, "--headless", "--convert-to", "pdf", "--outdir", out_dir, in_path])
        out_path = os.path.join(td, "input.pdf")
        with open(out_path, "rb") as f:
            data = f.read()
        return _bytesio(data)

def pdf_to_docx(file_bytes: bytes) -> io.BytesIO:
    # Try LibreOffice first, but fallback to text extraction if it fails
    try:
        soffice = _which(["soffice", "soffice.exe"], "libreoffice_path")
        if soffice:
            with tempfile.TemporaryDirectory() as td:
                in_path = os.path.join(td, "input.pdf")
                out_dir = td
                with open(in_path, "wb") as f:
                    f.write(file_bytes)
                
                # Try multiple conversion methods
                try:
                    # Method 1: Direct conversion with writer filter
                    _run_cmd([soffice, "--headless", "--convert-to", "docx", "--writer", "--outdir", out_dir, in_path], timeout=180)
                except RuntimeError:
                    try:
                        # Method 2: Using export filter
                        _run_cmd([soffice, "--headless", "--convert-to", "docx:MS Word 2007 XML", "--outdir", out_dir, in_path], timeout=180)
                    except RuntimeError:
                        try:
                            # Method 3: Generic conversion
                            _run_cmd([soffice, "--headless", "--convert-to", "docx", "--outdir", out_dir, in_path], timeout=180)
                        except RuntimeError:
                            raise RuntimeError("LibreOffice conversion failed")
                
                out_path = os.path.join(td, "input.docx")
                if not os.path.exists(out_path):
                    # Try alternative output filename
                    alt_path = os.path.join(td, os.path.splitext(os.path.basename(in_path))[0] + ".docx")
                    if os.path.exists(alt_path):
                        out_path = alt_path
                    else:
                        raise RuntimeError("LibreOffice conversion failed")
                
                with open(out_path, "rb") as f:
                    data = f.read()
                return _bytesio(data)
    except Exception:
        pass  # Fall back to text extraction method
    
    # Fallback: Use PyMuPDF to extract text and create a simple Word document
    try:
        from docx import Document
        from docx.shared import Inches
        
        doc = Document()
        doc.add_heading('PDF to Word Conversion', 0)
        
        # Extract text from PDF
        pdf_doc = fitz.open(stream=file_bytes, filetype="pdf")
        
        for page_num in range(len(pdf_doc)):
            page = pdf_doc.load_page(page_num)
            text = page.get_text()
            
            if text.strip():
                doc.add_heading(f'Page {page_num + 1}', level=1)
                doc.add_paragraph(text)
            else:
                doc.add_heading(f'Page {page_num + 1} (No text content)', level=1)
                doc.add_paragraph('This page contains no extractable text. It may contain images or scanned content.')
        
        # Save to BytesIO
        out = io.BytesIO()
        doc.save(out)
        out.seek(0)
        return out
        
    except ImportError:
        # If python-docx is not available, create a simple text file
        pdf_doc = fitz.open(stream=file_bytes, filetype="pdf")
        
        # Create RTF content (basic Word-compatible format)
        rtf_content = r"{\rtf1\ansi\deff0"
        rtf_content += r"{\fonttbl{\f0 Times New Roman;}}"
        rtf_content += r"\f0\fs24 "
        
        for page_num in range(len(pdf_doc)):
            page = pdf_doc.load_page(page_num)
            text = page.get_text()
            
            if text.strip():
                rtf_content += f"\\b Page {page_num + 1}\\b0\\par\\par "
                # Escape RTF special characters
                text = text.replace('\\', '\\\\').replace('{', '\\{').replace('}', '\\}')
                text = text.replace('\n', '\\par ')
                rtf_content += text + "\\par\\par "
            else:
                rtf_content += f"\\b Page {page_num + 1} (No text content)\\b0\\par\\par "
                rtf_content += "This page contains no extractable text. It may contain images or scanned content.\\par\\par "
        
        rtf_content += "}"
        
        return _bytesio(rtf_content.encode('utf-8'))
    
    except Exception as e:
        raise RuntimeError(f"PDF to Word conversion failed: {str(e)}. For best results, try using a dedicated PDF to Word converter or ensure your PDF contains extractable text.")

def pdf_to_xlsx(file_bytes: bytes) -> io.BytesIO:
    soffice = _which(["soffice", "soffice.exe"], "libreoffice_path")
    if not soffice:
        raise RuntimeError("LibreOffice not found. Please install LibreOffice from https://www.libreoffice.org/download/download-libreoffice/ and set the path in Settings, or ensure it's in your system PATH.")
    
    with tempfile.TemporaryDirectory() as td:
        in_path = os.path.join(td, "input.pdf")
        out_dir = td
        with open(in_path, "wb") as f:
            f.write(file_bytes)
        
        # Try multiple conversion methods
        try:
            # Method 1: Direct conversion with calc filter
            _run_cmd([soffice, "--headless", "--convert-to", "xlsx", "--calc", "--outdir", out_dir, in_path], timeout=180)
        except RuntimeError:
            try:
                # Method 2: Using export filter
                _run_cmd([soffice, "--headless", "--convert-to", "xlsx:MS Excel 2007 XML", "--outdir", out_dir, in_path], timeout=180)
            except RuntimeError:
                try:
                    # Method 3: Generic conversion
                    _run_cmd([soffice, "--headless", "--convert-to", "xlsx", "--outdir", out_dir, in_path], timeout=180)
                except RuntimeError:
                    raise RuntimeError("PDF to Excel conversion failed. This may require additional LibreOffice components. Try installing the complete LibreOffice suite with all optional components.")
        
        out_path = os.path.join(td, "input.xlsx")
        if not os.path.exists(out_path):
            # Try alternative output filename
            alt_path = os.path.join(td, os.path.splitext(os.path.basename(in_path))[0] + ".xlsx")
            if os.path.exists(alt_path):
                out_path = alt_path
            else:
                raise RuntimeError("Conversion completed but output file not found")
        
        with open(out_path, "rb") as f:
            data = f.read()
        return _bytesio(data)

def pdf_to_pptx(file_bytes: bytes) -> io.BytesIO:
    soffice = _which(["soffice", "soffice.exe"], "libreoffice_path")
    if not soffice:
        raise RuntimeError("LibreOffice not found. Please install LibreOffice from https://www.libreoffice.org/download/download-libreoffice/ and set the path in Settings, or ensure it's in your system PATH.")
    
    with tempfile.TemporaryDirectory() as td:
        in_path = os.path.join(td, "input.pdf")
        out_dir = td
        with open(in_path, "wb") as f:
            f.write(file_bytes)
        
        # Try multiple conversion methods
        try:
            # Method 1: Direct conversion with impress filter
            _run_cmd([soffice, "--headless", "--convert-to", "pptx", "--impress", "--outdir", out_dir, in_path], timeout=180)
        except RuntimeError:
            try:
                # Method 2: Using export filter
                _run_cmd([soffice, "--headless", "--convert-to", "pptx:MS PowerPoint 2007 XML", "--outdir", out_dir, in_path], timeout=180)
            except RuntimeError:
                try:
                    # Method 3: Generic conversion
                    _run_cmd([soffice, "--headless", "--convert-to", "pptx", "--outdir", out_dir, in_path], timeout=180)
                except RuntimeError:
                    raise RuntimeError("PDF to PowerPoint conversion failed. This may require additional LibreOffice components. Try installing the complete LibreOffice suite with all optional components.")
        
        out_path = os.path.join(td, "input.pptx")
        if not os.path.exists(out_path):
            # Try alternative output filename
            alt_path = os.path.join(td, os.path.splitext(os.path.basename(in_path))[0] + ".pptx")
            if os.path.exists(alt_path):
                out_path = alt_path
            else:
                raise RuntimeError("Conversion completed but output file not found")
        
        with open(out_path, "rb") as f:
            data = f.read()
        return _bytesio(data)

def ocr_pdf(file_bytes: bytes, dpi: int = 200) -> io.BytesIO:
    tesseract = _which(["tesseract", "tesseract.exe"], "tesseract_path")
    if not tesseract:
        raise RuntimeError("Tesseract not found")
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    writer = PdfWriter()
    for i in range(len(doc)):
        page = doc.load_page(i)
        zoom = dpi / 72.0
        mat = fitz.Matrix(zoom, zoom)
        pix = page.get_pixmap(matrix=mat, alpha=False)
        img_bytes = io.BytesIO(pix.tobytes("png"))
        pdf_bytes = pytesseract.image_to_pdf_or_hocr(Image.open(img_bytes), extension="pdf")
        r = PdfReader(io.BytesIO(pdf_bytes))
        writer.add_page(r.pages[0])
    out = io.BytesIO()
    writer.write(out)
    out.seek(0)
    return out

def compress_pdf_ghostscript(file_bytes: bytes, quality: str = "screen") -> io.BytesIO:
    gs = _which(["gswin64c", "gswin32c", "gs"], "ghostscript_path")
    if not gs:
        return _bytesio(file_bytes)
    with tempfile.TemporaryDirectory() as td:
        in_path = os.path.join(td, "in.pdf")
        out_path = os.path.join(td, "out.pdf")
        with open(in_path, "wb") as f:
            f.write(file_bytes)
        _run_cmd([gs, "-sDEVICE=pdfwrite", "-dCompatibilityLevel=1.4", f"-dPDFSETTINGS=/{quality}", "-dNOPAUSE", "-dQUIET", "-dBATCH", f"-sOutputFile={out_path}", in_path], cwd=td, timeout=180)
        with open(out_path, "rb") as f:
            data = f.read()
        return _bytesio(data)

def split_pdf_each_page(buffer: io.BytesIO) -> List[Tuple[str, bytes]]:
    reader = PdfReader(buffer)
    items = []
    for i, p in enumerate(reader.pages, start=1):
        w = PdfWriter()
        w.add_page(p)
        out = io.BytesIO()
        w.write(out)
        items.append((f"page_{i}.pdf", out.getvalue()))
    return items

def split_pdf_by_chunks(buffer: io.BytesIO, chunk_size: int) -> List[Tuple[str, bytes]]:
    reader = PdfReader(buffer)
    items = []
    n = len(reader.pages)
    size = max(1, chunk_size)
    idx = 0
    part = 1
    while idx < n:
        w = PdfWriter()
        for j in range(idx, min(idx + size, n)):
            w.add_page(reader.pages[j])
        out = io.BytesIO()
        w.write(out)
        items.append((f"part_{part}.pdf", out.getvalue()))
        part += 1
        idx += size
    return items

def reorder_pages(buffer: io.BytesIO, order: List[int]) -> io.BytesIO:
    reader = PdfReader(buffer)
    w = PdfWriter()
    n = len(reader.pages)
    for p in order:
        i = max(1, min(n, p)) - 1
        w.add_page(reader.pages[i])
    out = io.BytesIO()
    w.write(out)
    out.seek(0)
    return out

def remove_pages(buffer: io.BytesIO, remove_list: List[int]) -> io.BytesIO:
    reader = PdfReader(buffer)
    w = PdfWriter()
    n = len(reader.pages)
    remove_idx = {max(1, min(n, p)) - 1 for p in remove_list}
    for i, p in enumerate(reader.pages):
        if i in remove_idx:
            continue
        w.add_page(p)
    out = io.BytesIO()
    w.write(out)
    out.seek(0)
    return out

def extract_images_embedded(file_bytes: bytes) -> List[Tuple[str, bytes]]:
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    items = []
    for i in range(len(doc)):
        page = doc.load_page(i)
        imgs = page.get_images(full=True)
        k = 1
        for img in imgs:
            xref = img[0]
            data = doc.extract_image(xref)
            items.append((f"page_{i+1}_img_{k}.{data['ext']}", data["image"]))
            k += 1
    return items

def add_text_overlay(file_bytes: bytes, text: str, x: float, y: float, font_size: int, pages: Optional[List[int]]) -> io.BytesIO:
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    n = len(doc)
    targets = None
    if pages:
        targets = {max(1, min(n, p)) - 1 for p in pages}
    for i in range(n):
        if targets is None or i in targets:
            page = doc.load_page(i)
            page.insert_text((x, y), text, fontsize=font_size, fontname="helv")
    out = io.BytesIO(doc.tobytes())
    out.seek(0)
    return out

def add_image_overlay(file_bytes: bytes, image_bytes: bytes, x: float, y: float, width: float, height: float, pages: Optional[List[int]]) -> io.BytesIO:
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    n = len(doc)
    targets = None
    if pages:
        targets = {max(1, min(n, p)) - 1 for p in pages}
    for i in range(n):
        if targets is None or i in targets:
            page = doc.load_page(i)
            rect = fitz.Rect(x, y, x + width, y + height)
            page.insert_image(rect, stream=image_bytes)
    out = io.BytesIO(doc.tobytes())
    out.seek(0)
    return out

def _parse_rects(rects_str: str) -> List[Tuple[float, float, float, float]]:
    rects = []
    for item in rects_str.split(";"):
        item = item.strip()
        if not item:
            continue
        parts = [p.strip() for p in item.split(",")]
        if len(parts) != 4:
            continue
        x = float(parts[0]); y = float(parts[1]); w = float(parts[2]); h = float(parts[3])
        rects.append((x, y, x + w, y + h))
    return rects

def redact_pdf(file_bytes: bytes, rects_str: str, pages: Optional[List[int]]) -> io.BytesIO:
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    n = len(doc)
    targets = None
    if pages:
        targets = {max(1, min(n, p)) - 1 for p in pages}
    rects = _parse_rects(rects_str)
    for i in range(n):
        if targets is None or i in targets:
            page = doc.load_page(i)
            for (x1, y1, x2, y2) in rects:
                page.add_redact_annot(fitz.Rect(x1, y1, x2, y2))
            page.apply_redactions()
    out = io.BytesIO(doc.tobytes())
    out.seek(0)
    return out

def highlight_pdf(file_bytes: bytes, rects_str: str, pages: Optional[List[int]]) -> io.BytesIO:
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    n = len(doc)
    targets = None
    if pages:
        targets = {max(1, min(n, p)) - 1 for p in pages}
    rects = _parse_rects(rects_str)
    for i in range(n):
        if targets is None or i in targets:
            page = doc.load_page(i)
            for (x1, y1, x2, y2) in rects:
                page.add_highlight_annot(fitz.Rect(x1, y1, x2, y2))
    out = io.BytesIO(doc.tobytes())
    out.seek(0)
    return out

def edit_pdf_metadata(file_bytes: bytes, title: str = "", author: str = "", subject: str = "", keywords: str = "") -> io.BytesIO:
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    metadata = doc.metadata
    if title:
        metadata["title"] = title
    if author:
        metadata["author"] = author
    if subject:
        metadata["subject"] = subject
    if keywords:
        metadata["keywords"] = keywords
    doc.set_metadata(metadata)
    out = io.BytesIO(doc.tobytes())
    out.seek(0)
    return out

def get_pdf_metadata(file_bytes: bytes) -> dict:
    doc = fitz.open(stream=file_bytes, filetype="pdf")
    return doc.metadata

def compare_pdfs(file1_bytes: bytes, file2_bytes: bytes) -> io.BytesIO:
    doc1 = fitz.open(stream=file1_bytes, filetype="pdf")
    doc2 = fitz.open(stream=file2_bytes, filetype="pdf")
    
    writer = PdfWriter()
    
    # Add pages from first document
    for i, page in enumerate(doc1):
        writer.add_page(page)
    
    # Add separator page
    packet = io.BytesIO()
    c = canvas.Canvas(packet, pagesize=letter)
    c.setFont("Helvetica", 24)
    c.setFillColor(black)
    c.drawCentredString(letter[0]/2, letter[1]/2, "--- COMPARISON SEPARATOR ---")
    c.save()
    packet.seek(0)
    separator = PdfReader(packet).pages[0]
    writer.add_page(separator)
    
    # Add pages from second document
    for page in doc2:
        writer.add_page(page)
    
    out = io.BytesIO()
    writer.write(out)
    out.seek(0)
    return out
